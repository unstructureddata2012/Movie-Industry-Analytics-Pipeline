import csv
import xml.etree.ElementTree as ET
import os
import json
import sys
import pdfplumber
import re
from docx import Document
from openpyxl import load_workbook
from datetime import datetime

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'src')))
from storage.mongo import save_to_mongo 

def extract_movie_fields(movie):
    return {
        "id": movie.get("id"),
        "title": movie.get("title"),
        "release_date": movie.get("release_date"),
        "popularity": movie.get("popularity")
    }

def parse_csv_file(file_path):
    with open(file_path, "r") as f:
        reader = csv.DictReader(f)  
        for row in reader:
            # print(f"ID: {row['id']}, Title: {row['title']}, Release Date: {row['release_date']}, Popularity: {row['popularity']}")
            movie_fields = extract_movie_fields(row)  
            print(f"Saving to MongoDB: {movie_fields}")
            save_to_mongo(movie_fields, "CSV Source")
            

def parse_xml_file(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    for item in root.findall("movie"):
        movie = {
            "id": item.find("id").text,
            "title": item.find("title").text,
            "release_date": item.find("release_date").text,
            "popularity": item.find("popularity").text
        }
        # print(f"ID: {movie['id']}, Title: {movie['title']}, Release Date: {movie['release_date']}, Popularity: {movie['popularity']}")
        movie_fields = extract_movie_fields(movie) 
        print(f"Saving to MongoDB: {movie_fields}")
        save_to_mongo(movie_fields, "XML Source")

def parse_json_files():
    folder_path = "../../data/raw/api/"

    for filename in os.listdir(folder_path):
        if filename.endswith(".json"): 
            file_path = os.path.join(folder_path, filename)

            with open(file_path, "r") as f:
                movie_data = json.load(f)

            if isinstance(movie_data, dict):  # If movie_data is a dictionary (single movie)
                movie_fields = extract_movie_fields(movie_data)
                # print(movie_fields)
                print(f"Saving to MongoDB: {movie_fields}")
                save_to_mongo(movie_fields, filename)
            else:
                print(f"Unexpected structure in {filename}: {movie_data}")


def extract_text_from_pdf(pdf_path):
    pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(normalize_text(text))

    return "\n\n".join(pages)

def normalize_text(text):
    if not text:
        return ""

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()

import chardet


def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    confidence = result.get("confidence")

    print(f"Detected encoding: {encoding} (confidence: {confidence})")

    text = raw.decode(encoding, errors="replace")

    return text

def extract_text_from_two_column_pdf(pdf_path, gap=10):
    pages_text = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            mid_x = page.width / 2

            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))

            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")

            combined = "\n\n".join(part for part in [left_text, right_text] if part)
            if combined:
                pages_text.append(combined)

    return "\n\n".join(pages_text)

def extract_text_from_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

def extract_text_from_two_column_word(docx_path):
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)

def extract_data_from_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb["Movie Data"]

    movies = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        movie = {
            "id": row[0],
            "title": row[1],
            "genre": row[2],
            "release_year": row[3],
            "director": row[4],
            "budget_usd": row[5],
            "revenue_usd": row[6],
            "rating_imdb": row[7],
            "country": row[8]
        }
        movies.append(movie)

    return movies
def extract_summary_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Summary"]

    summary = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        metric = row[0]
        value = row[1]
        summary[metric] = value

    return summary
if __name__ == "__main__":
    # parse_json_files()
    # parse_csv_file("../../data/raw/csv/sample.csv")
    # parse_xml_file("../../data/raw/xml/sample.xml")

    # pdf_path = "../../data/raw/pdf/film_two_column.pdf"
    text_path = "../../data/raw/word/film_two_column.docx"
    # text = extract_text_from_pdf(pdf_path)
    # text = extract_text_from_two_column_pdf(pdf_path)
    # text = extract_text_from_two_column_word(text_path)
    # print(text)
    excel_path = "../../data/raw/excel/movies_excel.xlsx"

    movies = extract_data_from_excel(excel_path)
    print(movies)

    summary = extract_summary_from_excel(excel_path)
    print(summary)


