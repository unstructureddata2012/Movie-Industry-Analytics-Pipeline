#!/usr/bin/env python3
"""
generate_lab12_doc.py
=====================
Generates the Lab 12 Word document:
    docs/labs/lab12_data_visualization.docx

Run from the project root:
    python scripts/generate_lab12_doc.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ───────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x6F, 0xAF)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def code_block(doc, code_text):
    """Monospace shaded paragraph for code snippets."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
            row.cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
    return table


# ── document ──────────────────────────────────────────────────────────────────

def build_document(out_path: Path):
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("International Burch University")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x6F, 0xAF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Faculty of Engineering and Natural Sciences")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = course.add_run("IT 2012 – Unstructured Data\nAcademic Year 2024/2025")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()
    lab_title = doc.add_paragraph()
    lab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = lab_title.add_run("LAB 12")
    r3.font.size = Pt(28)
    r3.bold = True
    r3.font.color.rgb = RGBColor(0x1A, 0x6F, 0xAF)

    lab_sub = doc.add_paragraph()
    lab_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = lab_sub.add_run("Data Visualization")
    r4.font.size = Pt(22)
    r4.bold = True
    r4.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    doc.add_paragraph()
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = meta.add_run(
        "Student: Amila Causevic\n"
        "E-mail: amila.causevic@stu.ibu.edu.ba\n"
        "Instructor: Assist. Prof. Dr. Dželila Mehanović\n"
        "Date: May 2025"
    )
    r5.font.size = Pt(12)
    r5.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION & OBJECTIVES
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "1. Introduction and Lab Objectives", level=1)
    body(doc,
        "Data visualization is the process of representing data graphically "
        "so that patterns, trends, and insights can be perceived immediately. "
        "Effective visualization bridges raw data and human understanding—it "
        "translates numbers into stories. This lab implements a full-stack "
        "visualization module for the Movie Industry Analytics Pipeline, "
        "producing both publication-ready static figures and interactive "
        "HTML-based exploratory charts.")

    heading(doc, "Objectives", level=2)
    for obj in [
        "Understand the matplotlib object-oriented API (Figure / Axes model) "
        "and use it to construct multi-panel layouts.",
        "Apply seaborn to generate statistical visualizations (histograms, "
        "boxplots, scatter plots, heatmaps) with minimal boilerplate.",
        "Use Plotly Express and Plotly Graph Objects to build interactive "
        "charts with rich hover tooltips, saved as self-contained HTML files.",
        "Automate chart generation through a reusable Python module and "
        "CLI entry-point script.",
        "Apply Tufte's principles of graphical excellence: maximize "
        "data-ink ratio, avoid chartjunk, choose chart types that match "
        "the question being asked.",
        "Document visualization choices explaining why each chart type was "
        "selected for the underlying data characteristic.",
    ]:
        bullet(doc, obj)

    doc.add_paragraph()

    # ── Lab structure ──
    heading(doc, "Lab Structure", level=2)
    body(doc,
        "This lab introduces a new src/visualization/ package with dedicated "
        "modules for static and interactive chart generation. It also adds a "
        "CLI script scripts/generate_visualizations.py, a notebook "
        "notebook/lab12_visualization.ipynb, and two output directories under "
        "outputs/visualizations/. The pipeline entry point run_pipeline.py is "
        "updated with a new run_visualizations_pipeline() step.")

    body(doc, "In src/visualization/ you need to add the following:")

    structure_items = [
        ("src/visualization/__init__.py",
         "makes the visualization folder a Python package and re-exports all "
         "chart functions so they can be imported from a single location."),
        ("src/visualization/static_charts.py",
         "contains 8 matplotlib and seaborn chart functions. Each function "
         "accepts a DataFrame and an output directory, draws a figure using "
         "the object-oriented API, and saves it as both PNG (300 dpi) and PDF."),
        ("src/visualization/interactive_charts.py",
         "contains 5 Plotly Express and Graph Objects chart functions. Each "
         "function saves a self-contained interactive HTML file that can be "
         "opened in any browser without a running server."),
        ("src/visualization/chart_generator.py",
         "is the orchestrator module. It loads the cleaned dataset, iterates "
         "over all static and interactive chart functions, and calls each one "
         "in sequence. This is the module imported by the CLI script."),
    ]

    for name, desc in structure_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run_bold = p.add_run(name + "  ")
        run_bold.bold = True
        p.add_run("– " + desc)

    doc.add_paragraph()
    body(doc, "You also need to create the following at the project root level:")

    root_items = [
        ("scripts/generate_visualizations.py",
         "is the CLI entry-point. Run it from the project root with "
         "python scripts/generate_visualizations.py to regenerate all charts. "
         "It accepts an optional --data flag to point at a different CSV file."),
        ("notebook/lab12_visualization.ipynb",
         "is the main submission notebook demonstrating all lab steps "
         "interactively: setup, each static chart, each interactive chart, "
         "the automated generator, and the Assignment 12 solution."),
    ]

    for name, desc in root_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        run_bold = p.add_run(name + "  ")
        run_bold.bold = True
        p.add_run("– " + desc)

    doc.add_paragraph()
    body(doc,
        "Inside outputs/, two new folders store all generated chart files. "
        "outputs/visualizations/static/ holds the PNG and PDF exports produced "
        "by matplotlib and seaborn — you do not need to commit these files to "
        "version control as they are fully reproducible by running the script. "
        "outputs/visualizations/interactive/ holds the HTML files produced by "
        "Plotly — open any of them directly in a browser to interact with the chart.")

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 1 – INSTALL
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 1 – Install Required Libraries and Tools", level=1)
    body(doc,
        "Open requirements.txt in the project root and add the following lines "
        "if they are not already present:")

    code_block(doc,
        "matplotlib>=3.8.0\n"
        "seaborn>=0.12.0\n"
        "plotly>=6.0.0\n"
        "kaleido>=0.2.1"
    )
    body(doc,
        "Then activate your virtual environment in the VS Code terminal and "
        "install everything:")
    code_block(doc, "pip install -r requirements.txt")
    body(doc,
        "If you only want the visualization packages for this lab:")
    code_block(doc,
        "pip install matplotlib seaborn plotly kaleido"
    )
    body(doc,
        "Note: kaleido is required only if you want to export Plotly figures as "
        "static PNG or PDF files. For HTML-only output it is not needed.")

    heading(doc, "1.1 Setting Up a Virtual Environment", level=2)
    body(doc,
        "If you have not already created a virtual environment, do it once from "
        "the project root:")
    code_block(doc,
        "# macOS / Linux\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n\n"
        "# Windows (Command Prompt)\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate.bat"
    )
    body(doc,
        "When activated your terminal prompt will be prefixed with (.venv). "
        "Always activate the environment before running any lab code.")

    heading(doc, "1.2 Verifying the Installation", level=2)
    body(doc,
        "After installation, confirm everything imported correctly:")
    code_block(doc,
        "python -c \"import matplotlib, seaborn, plotly, pandas, numpy; \\\n"
        "print('matplotlib', matplotlib.__version__); \\\n"
        "print('seaborn',    seaborn.__version__); \\\n"
        "print('plotly',     plotly.__version__); \\\n"
        "print('pandas',     pandas.__version__); \\\n"
        "print('numpy',      numpy.__version__)\""
    )
    body(doc, "Expected output (your versions may be higher, that is fine):")
    code_block(doc,
        "matplotlib 3.10.8\n"
        "seaborn    0.13.2\n"
        "plotly     6.7.0\n"
        "pandas     2.2.3\n"
        "numpy      1.26.4"
    )
    body(doc,
        "If you see an ImportError for any package, re-run the pip install "
        "command above and make sure the virtual environment is activated.")

    heading(doc, "1.3 Opening the Lab Notebook", level=2)
    body(doc,
        "This lab includes a Jupyter notebook at notebook/lab12_visualization.ipynb. "
        "Launch JupyterLab from the project root (not from inside the notebook/ folder) "
        "so that the relative paths used in the notebook resolve correctly:")
    code_block(doc,
        "jupyter lab\n"
        "# or\n"
        "jupyter notebook"
    )
    body(doc,
        "Open notebook/lab12_visualization.ipynb and run cells top to bottom "
        "using Shift+Enter.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 2 – UNDERSTANDING DATA VISUALIZATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 2 – Understanding Data Visualization", level=1)
    body(doc,
        "A visualization communicates a data insight faster than a table of "
        "numbers ever could. Before writing any code, the most important decision "
        "is which chart type matches the question you are asking. The table below "
        "shows the mapping used throughout this lab:")

    add_table(doc,
        ["Question you want to answer", "Chart type", "Library"],
        [
            ["Which movies earned the most?",           "Horizontal bar chart",     "matplotlib"],
            ["How has the average rating changed?",     "Dual-axis line + bar",     "matplotlib"],
            ["Does higher budget mean higher revenue?", "Scatter plot",             "seaborn"],
            ["What does the rating distribution look like?", "Histogram + KDE",    "seaborn"],
            ["How do ratings vary across genres?",      "Box-and-whisker plot",     "seaborn"],
            ["What correlates with what?",              "Heatmap",                  "seaborn"],
            ["How many movies per genre?",              "Vertical bar chart",       "matplotlib"],
            ["All key views at once",                   "2×2 multi-panel layout",   "matplotlib"],
        ]
    )
    doc.add_paragraph()

    heading(doc, "2.1 matplotlib – Object-Oriented API", level=2)
    body(doc,
        "matplotlib is Python's foundational plotting library. This lab uses "
        "the object-oriented API exclusively. Every chart starts with an explicit "
        "Figure and Axes object so every visual element can be configured precisely:")
    code_block(doc,
        "import matplotlib\n"
        "matplotlib.use('Agg')  # non-interactive backend for scripts\n"
        "import matplotlib.pyplot as plt\n\n"
        "fig, ax = plt.subplots(figsize=(10, 6))\n"
        "ax.barh(titles, values)\n"
        "ax.set_xlabel('Revenue (USD Billion)')\n"
        "fig.tight_layout()\n"
        "fig.savefig('chart.png', dpi=300, bbox_inches='tight')\n"
        "fig.savefig('chart.pdf', bbox_inches='tight')"
    )
    body(doc,
        "The Agg backend must be set before importing pyplot. It allows "
        "savefig() to work in scripts and notebooks without opening a display window.")

    heading(doc, "2.2 seaborn – Statistical Charts", level=2)
    body(doc,
        "seaborn is built on top of matplotlib and provides ready-made statistical "
        "chart functions (histplot, boxplot, scatterplot, heatmap) with much less "
        "boilerplate. Always call sns.set_theme() once at the top of the notebook "
        "to apply consistent styling:")
    code_block(doc,
        "import seaborn as sns\n\n"
        "sns.set_theme(style='whitegrid')\n"
        "sns.set_context('notebook')\n"
        "sns.set_palette('viridis')"
    )
    body(doc,
        "Note: since seaborn 0.12, passing a palette without a hue= parameter is "
        "deprecated. For boxplots and similar charts, always pass "
        "hue='column_name' with legend=False when you want colour by category.")

    heading(doc, "2.3 Plotly – Interactive Charts", level=2)
    body(doc,
        "Plotly produces interactive web charts saved as self-contained HTML files. "
        "It has two APIs: plotly.express (px) for one-line charts with hover, zoom, "
        "and legend toggle; and plotly.graph_objects (go) for lower-level control "
        "and multi-panel layouts with make_subplots(). Every chart is saved with "
        "fig.write_html() and can be opened in any browser without a server:")
    code_block(doc,
        "import plotly.express as px\n\n"
        "fig = px.scatter(data, x='budget_M', y='revenue_M',\n"
        "                 color='primary_genre', hover_name='title',\n"
        "                 template='plotly_white')\n"
        "fig.write_html('outputs/visualizations/interactive/chart.html')"
    )

    heading(doc, "2.4 Setting Up the Package", level=2)
    body(doc,
        "In src/visualization/ you need to add the __init__.py file first. "
        "This file makes the visualization folder a Python package so its modules "
        "can be imported in other parts of the project. Create the file with the "
        "following content that re-exports all chart functions from a single location:")
    code_block(doc,
        "# src/visualization/__init__.py\n"
        "from .static_charts import (\n"
        "    plot_top_movies_by_revenue,\n"
        "    plot_avg_rating_over_years,\n"
        "    plot_budget_vs_revenue_scatter,\n"
        "    plot_rating_distribution,\n"
        "    plot_genre_rating_boxplot,\n"
        "    plot_correlation_heatmap,\n"
        "    plot_genre_count_bar,\n"
        "    plot_dashboard_subplots,\n"
        ")\n"
        "from .interactive_charts import (\n"
        "    interactive_budget_vs_revenue,\n"
        "    interactive_top_movies_bar,\n"
        "    interactive_movies_per_year,\n"
        "    interactive_genre_boxplot,\n"
        "    interactive_multi_layout,\n"
        ")"
    )
    body(doc,
        "After that, open notebook/lab12_visualization.ipynb and run the first "
        "setup cell. It adds src/ to the Python path, sets the Agg backend, "
        "imports all libraries, creates the output directories, and loads the "
        "cleaned dataset from data/processed/cleaned/movies_clean.csv.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 3 – BAR AND LINE CHARTS WITH MATPLOTLIB
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 3 – Static Bar and Line Charts with Matplotlib", level=1)
    body(doc,
        "In this part you need to create the file src/visualization/static_charts.py "
        "and add the first two chart functions to it. This module will hold all "
        "eight static chart functions used in the lab. Start the file with the "
        "following imports and a shared _save() helper that every chart function "
        "will call to write PNG and PDF files:")
    code_block(doc,
        "import matplotlib\n"
        "matplotlib.use('Agg')\n"
        "import matplotlib.pyplot as plt\n"
        "import matplotlib.ticker as mticker\n"
        "import seaborn as sns\n"
        "import pandas as pd\n"
        "import numpy as np\n"
        "from pathlib import Path\n\n"
        "sns.set_theme(style='whitegrid')\n"
        "sns.set_context('notebook')\n"
        "sns.set_palette('viridis')\n\n"
        "STATIC_OUT = Path('outputs/visualizations/static')\n\n"
        "def _save(fig, stem, out_dir=STATIC_OUT):\n"
        "    out_dir.mkdir(parents=True, exist_ok=True)\n"
        "    png_path = out_dir / f'{stem}.png'\n"
        "    pdf_path = out_dir / f'{stem}.pdf'\n"
        "    fig.savefig(png_path, dpi=300, bbox_inches='tight')\n"
        "    fig.savefig(pdf_path, bbox_inches='tight')\n"
        "    plt.close(fig)\n"
        "    return {'png': str(png_path), 'pdf': str(pdf_path)}"
    )

    heading(doc, "3.1 Top 10 Movies by Revenue – Horizontal Bar Chart", level=2)
    body(doc,
        "Add the first function to static_charts.py. A horizontal bar chart is "
        "used here because the category labels (movie titles) are long strings "
        "that read more naturally from left to right. Revenue values are "
        "converted to USD billions and inline annotations are added so the "
        "reader does not have to look up each bar on the x-axis:")
    code_block(doc,
        "def plot_top_movies_by_revenue(df, n=10, out_dir=STATIC_OUT):\n"
        "    top = (df[df['revenue_usd'] > 0]\n"
        "           .nlargest(n, 'revenue_usd')[['title', 'revenue_usd']]\n"
        "           .sort_values('revenue_usd'))\n\n"
        "    fig, ax = plt.subplots(figsize=(10, 6))\n"
        "    colors = sns.color_palette('viridis', n)\n"
        "    bars = ax.barh(top['title'], top['revenue_usd'] / 1e9, color=colors)\n\n"
        "    for bar in bars:\n"
        "        ax.text(bar.get_width() + 0.03,\n"
        "                bar.get_y() + bar.get_height() / 2,\n"
        "                f'${bar.get_width():.2f}B', va='center', fontsize=9)\n\n"
        "    ax.set_xlabel('Box-Office Revenue (USD Billion)')\n"
        "    ax.set_title(f'Top {n} Movies by Revenue', fontweight='bold')\n"
        "    ax.xaxis.set_major_formatter(\n"
        "        mticker.FuncFormatter(lambda x, _: f'${x:.1f}B'))\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'top_movies_by_revenue', out_dir)"
    )
    body(doc,
        "In the notebook, call the function directly and run plt.show() to "
        "preview the figure inline before saving.")

    heading(doc, "3.2 Average Rating Over Years – Dual-Axis Line Chart", level=2)
    body(doc,
        "Add a second function to static_charts.py. This chart overlays two "
        "related but differently scaled series on the same x-axis: the number "
        "of movies released per year (bars, left y-axis) and the mean audience "
        "rating per year (line, right y-axis). Use ax.twinx() to create the "
        "second y-axis that shares the x-axis:")
    code_block(doc,
        "def plot_avg_rating_over_years(df, out_dir=STATIC_OUT):\n"
        "    yearly = (df.groupby('release_year')\n"
        "               .agg(avg_rating=('vote_average', 'mean'),\n"
        "                    movie_count=('title', 'count'))\n"
        "               .reset_index()\n"
        "               .query('release_year >= 1980 and release_year <= 2025'))\n\n"
        "    fig, ax1 = plt.subplots(figsize=(12, 5))\n"
        "    ax1.bar(yearly['release_year'], yearly['movie_count'],\n"
        "            color='#a8d5e2', alpha=0.5, label='Movie Count')\n"
        "    ax1.set_ylabel('Number of Movies', color='#a8d5e2')\n\n"
        "    ax2 = ax1.twinx()\n"
        "    ax2.plot(yearly['release_year'], yearly['avg_rating'],\n"
        "             color='#1a6faf', linewidth=2.5, marker='o', markersize=5)\n"
        "    ax2.set_ylabel('Average Vote Rating (0–10)', color='#1a6faf')\n"
        "    ax2.set_ylim(0, 10)\n\n"
        "    ax1.set_xlabel('Release Year')\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'avg_rating_over_years', out_dir)"
    )
    body(doc,
        "In the notebook, demonstrate this chart after the bar chart. "
        "In my output, the chart shows that the number of movies in the dataset "
        "increases towards recent years and that the average rating stays "
        "roughly between 6 and 8 throughout the period.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 4 – STATISTICAL CHARTS WITH SEABORN
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 4 – Statistical Charts with Seaborn", level=1)
    body(doc,
        "Continue adding functions to src/visualization/static_charts.py. "
        "This part uses seaborn's statistical chart functions. All functions "
        "follow the same pattern: create a fig, ax with plt.subplots(), pass "
        "ax= to the seaborn function, configure labels and title, then call _save().")

    heading(doc, "4.1 Budget vs Revenue – Scatter Plot", level=2)
    body(doc,
        "A scatter plot is the natural choice to show the relationship between "
        "two continuous variables. This function encodes genre as hue and "
        "vote_average as marker size, so three dimensions are visible at once. "
        "It also draws a diagonal break-even line to make profitability "
        "immediately visible:")
    code_block(doc,
        "def plot_budget_vs_revenue_scatter(df, out_dir=STATIC_OUT):\n"
        "    data = df[(df['budget_usd'] > 0) & (df['revenue_usd'] > 0)].copy()\n"
        "    data['budget_M']  = data['budget_usd']  / 1e6\n"
        "    data['revenue_M'] = data['revenue_usd'] / 1e6\n"
        "    genres = data['primary_genre'].value_counts().nlargest(6).index\n"
        "    data['genre_label'] = data['primary_genre'].where(\n"
        "        data['primary_genre'].isin(genres), other='Other')\n\n"
        "    fig, ax = plt.subplots(figsize=(11, 7))\n"
        "    sns.scatterplot(data=data, x='budget_M', y='revenue_M',\n"
        "                    hue='genre_label', palette='viridis',\n"
        "                    size='vote_average', sizes=(30, 250),\n"
        "                    alpha=0.75, ax=ax)\n"
        "    max_v = max(data['budget_M'].max(), data['revenue_M'].max()) * 1.05\n"
        "    ax.plot([0, max_v], [0, max_v], '--', color='gray',\n"
        "            linewidth=1.2, label='Break-even (1:1)')\n"
        "    ax.set_xlabel('Production Budget (USD Million)')\n"
        "    ax.set_ylabel('Box-Office Revenue (USD Million)')\n"
        "    ax.set_title('Budget vs Revenue – Coloured by Genre', fontweight='bold')\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'budget_vs_revenue_scatter', out_dir)"
    )

    heading(doc, "4.2 Rating Distribution – Histogram + KDE", level=2)
    body(doc,
        "Add a histogram with a KDE (kernel density estimate) overlay to show "
        "the full distribution shape of audience ratings. Add vertical reference "
        "lines for the mean and median so the reader can immediately see whether "
        "the distribution is skewed:")
    code_block(doc,
        "def plot_rating_distribution(df, out_dir=STATIC_OUT):\n"
        "    fig, ax = plt.subplots(figsize=(9, 5))\n"
        "    sns.histplot(df['vote_average'].dropna(), bins=20, kde=True,\n"
        "                 color='#1a6faf', edgecolor='white', ax=ax)\n"
        "    ax.axvline(df['vote_average'].mean(), color='firebrick',\n"
        "               linestyle='--', linewidth=1.8,\n"
        "               label=f\"Mean = {df['vote_average'].mean():.2f}\")\n"
        "    ax.axvline(df['vote_average'].median(), color='darkorange',\n"
        "               linestyle='-.', linewidth=1.8,\n"
        "               label=f\"Median = {df['vote_average'].median():.2f}\")\n"
        "    ax.set_xlabel('Vote Average (0–10)')\n"
        "    ax.set_ylabel('Number of Movies')\n"
        "    ax.set_title('Distribution of Movie Ratings', fontweight='bold')\n"
        "    ax.legend()\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'rating_distribution', out_dir)"
    )

    heading(doc, "4.3 Genre Rating Box Plot", level=2)
    body(doc,
        "A box plot shows the full distribution spread (median, quartiles, "
        "outliers) for each genre in a compact side-by-side layout. Sort genres "
        "by median rating descending so the highest-rated genre appears first. "
        "Since seaborn 0.12, always pass hue= and legend=False together when "
        "colouring by the same column as x=:")
    code_block(doc,
        "def plot_genre_rating_boxplot(df, out_dir=STATIC_OUT):\n"
        "    order = (df.groupby('primary_genre')['vote_average']\n"
        "               .median().sort_values(ascending=False).index.tolist())\n"
        "    fig, ax = plt.subplots(figsize=(12, 6))\n"
        "    sns.boxplot(data=df, x='primary_genre', y='vote_average',\n"
        "                order=order, hue='primary_genre',\n"
        "                palette='viridis', legend=False, ax=ax)\n"
        "    ax.set_xlabel('Genre')\n"
        "    ax.set_ylabel('Vote Average (0–10)')\n"
        "    ax.set_title('Movie Rating Distribution by Genre', fontweight='bold')\n"
        "    ax.tick_params(axis='x', rotation=30)\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'genre_rating_boxplot', out_dir)"
    )

    heading(doc, "4.4 Correlation Heatmap", level=2)
    body(doc,
        "A heatmap is the standard way to show an N×N correlation matrix. "
        "Use the coolwarm diverging palette so negative correlations appear blue "
        "and positive ones appear red, with white at zero. Set square=True so "
        "every cell has equal visual weight:")
    code_block(doc,
        "def plot_correlation_heatmap(df, out_dir=STATIC_OUT):\n"
        "    numeric_cols = ['budget_usd', 'revenue_usd', 'vote_average',\n"
        "                    'vote_count', 'popularity']\n"
        "    corr = df[numeric_cols].corr()\n"
        "    labels = {'budget_usd': 'Budget', 'revenue_usd': 'Revenue',\n"
        "              'vote_average': 'Avg Rating', 'vote_count': 'Vote Count',\n"
        "              'popularity': 'Popularity'}\n"
        "    corr.index = corr.columns = [labels[c] for c in corr.columns]\n\n"
        "    fig, ax = plt.subplots(figsize=(8, 6))\n"
        "    sns.heatmap(corr, annot=True, fmt='.2f', cmap='coolwarm',\n"
        "                center=0, vmin=-1, vmax=1, square=True,\n"
        "                linewidths=0.5, ax=ax)\n"
        "    ax.set_title('Correlation Matrix', fontweight='bold')\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'correlation_heatmap', out_dir)"
    )

    heading(doc, "4.5 Genre Count Bar Chart", level=2)
    body(doc,
        "Add a simple vertical bar chart to show how many movies belong to "
        "each primary genre. Add count labels above each bar so the reader "
        "does not have to read off the y-axis:")
    code_block(doc,
        "def plot_genre_count_bar(df, out_dir=STATIC_OUT):\n"
        "    counts = df['primary_genre'].value_counts()\n"
        "    fig, ax = plt.subplots(figsize=(10, 5))\n"
        "    bars = ax.bar(counts.index, counts.values,\n"
        "                  color=sns.color_palette('viridis', len(counts)))\n"
        "    for bar in bars:\n"
        "        ax.text(bar.get_x() + bar.get_width() / 2,\n"
        "                bar.get_height() + 0.5,\n"
        "                str(int(bar.get_height())),\n"
        "                ha='center', va='bottom', fontsize=9)\n"
        "    ax.set_xlabel('Primary Genre')\n"
        "    ax.set_ylabel('Number of Movies')\n"
        "    ax.set_title('Number of Movies per Genre', fontweight='bold')\n"
        "    ax.tick_params(axis='x', rotation=30)\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'genre_count_bar', out_dir)"
    )
    body(doc,
        "In the notebook, demonstrate each seaborn chart in its own cell. "
        "In my output, the scatter plot shows that most movies above the "
        "break-even line are Action or Adventure titles. The rating distribution "
        "is approximately bell-shaped around a mean of 7.0. The correlation "
        "heatmap shows that revenue and budget are the most strongly correlated "
        "pair (r ≈ 0.74), while vote_average has weak correlations with all "
        "financial variables.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 5 – MULTI-PANEL DASHBOARD
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 5 – Multi-Panel Dashboard Layout", level=1)
    body(doc,
        "Add the final function to src/visualization/static_charts.py. "
        "This function creates a 2×2 subplot figure that combines the four "
        "most informative views on a single canvas. Use plt.subplots(2, 2) "
        "to create the grid and address each panel by its row-column index. "
        "This is a common pattern for executive dashboards where a single "
        "shareable image must tell the whole story:")
    code_block(doc,
        "def plot_dashboard_subplots(df, out_dir=STATIC_OUT):\n"
        "    fig, axes = plt.subplots(2, 2, figsize=(14, 10))\n"
        "    fig.suptitle('Movie Industry Analytics Dashboard',\n"
        "                 fontsize=16, fontweight='bold', y=1.01)\n\n"
        "    # Panel A – top 10 revenue (top-left)\n"
        "    top = (df[df['revenue_usd'] > 0].nlargest(10, 'revenue_usd')\n"
        "           [['title', 'revenue_usd']].sort_values('revenue_usd'))\n"
        "    axes[0, 0].barh(top['title'], top['revenue_usd'] / 1e9,\n"
        "                    color=sns.color_palette('viridis', 10))\n"
        "    axes[0, 0].set_title('Top 10 by Revenue', fontweight='bold')\n\n"
        "    # Panel B – rating histogram (top-right)\n"
        "    sns.histplot(df['vote_average'].dropna(), bins=15, kde=True,\n"
        "                 color='#1a6faf', ax=axes[0, 1])\n"
        "    axes[0, 1].set_title('Rating Distribution', fontweight='bold')\n\n"
        "    # Panel C – genre count (bottom-left)\n"
        "    counts = df['primary_genre'].value_counts().head(8)\n"
        "    axes[1, 0].bar(counts.index, counts.values,\n"
        "                   color=sns.color_palette('viridis', len(counts)))\n"
        "    axes[1, 0].set_title('Movies per Genre (top 8)', fontweight='bold')\n"
        "    axes[1, 0].tick_params(axis='x', rotation=30)\n\n"
        "    # Panel D – budget vs revenue scatter (bottom-right)\n"
        "    data = df[(df['budget_usd'] > 0) & (df['revenue_usd'] > 0)].copy()\n"
        "    sc = axes[1, 1].scatter(\n"
        "        data['budget_usd'] / 1e6, data['revenue_usd'] / 1e6,\n"
        "        c=data['vote_average'], cmap='viridis', alpha=0.65, s=40)\n"
        "    fig.colorbar(sc, ax=axes[1, 1], label='Rating')\n"
        "    axes[1, 1].set_title('Budget vs Revenue', fontweight='bold')\n\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'dashboard_subplots', out_dir)"
    )
    body(doc,
        "In the notebook, call plot_dashboard_subplots(df) and display the "
        "result inline. You should see four panels on one canvas. "
        "After this, static_charts.py is complete with 8 functions. "
        "You can run a quick test by importing and calling each function "
        "individually in the notebook to confirm all 8 PNG and PDF files "
        "are created under outputs/visualizations/static/.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 6 – INTERACTIVE CHARTS WITH PLOTLY EXPRESS
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 6 – Interactive Charts with Plotly Express", level=1)
    body(doc,
        "Now create a new file src/visualization/interactive_charts.py. "
        "This module holds all five interactive chart functions. "
        "Start with the imports and a shared _save_html() helper:")
    code_block(doc,
        "import plotly.express as px\n"
        "import plotly.graph_objects as go\n"
        "from plotly.subplots import make_subplots\n"
        "import pandas as pd\n"
        "from pathlib import Path\n\n"
        "INTERACTIVE_OUT = Path('outputs/visualizations/interactive')\n"
        "TEMPLATE = 'plotly_white'\n\n"
        "def _save_html(fig, stem, out_dir=INTERACTIVE_OUT):\n"
        "    out_dir.mkdir(parents=True, exist_ok=True)\n"
        "    path = out_dir / f'{stem}.html'\n"
        "    fig.write_html(str(path))\n"
        "    return str(path)"
    )

    heading(doc, "6.1 Budget vs Revenue – Interactive Scatter", level=2)
    body(doc,
        "Add the first function. Use px.scatter() with color, size, hover_name, "
        "and hover_data to create a fully interactive chart. Add a break-even "
        "reference line using a go.Scatter trace with mode='lines':")
    code_block(doc,
        "def interactive_budget_vs_revenue(df, out_dir=INTERACTIVE_OUT):\n"
        "    data = df[(df['budget_usd'] > 0) & (df['revenue_usd'] > 0)].copy()\n"
        "    data['budget_M']  = (data['budget_usd']  / 1e6).round(1)\n"
        "    data['revenue_M'] = (data['revenue_usd'] / 1e6).round(1)\n"
        "    data['roi'] = ((data['revenue_usd'] - data['budget_usd'])\n"
        "                   / data['budget_usd'] * 100).round(1)\n\n"
        "    fig = px.scatter(\n"
        "        data, x='budget_M', y='revenue_M',\n"
        "        color='primary_genre', size='vote_count',\n"
        "        hover_name='title',\n"
        "        hover_data={'release_year': True,\n"
        "                    'vote_average': ':.2f', 'roi': ':.1f'},\n"
        "        title='Budget vs Box-Office Revenue – Interactive Explorer',\n"
        "        template=TEMPLATE,\n"
        "    )\n"
        "    max_v = max(data['budget_M'].max(), data['revenue_M'].max()) * 1.05\n"
        "    fig.add_trace(go.Scatter(\n"
        "        x=[0, max_v], y=[0, max_v], mode='lines',\n"
        "        line=dict(dash='dash', color='gray', width=1),\n"
        "        name='Break-even', hoverinfo='skip'))\n"
        "    fig.update_layout(legend_title='Genre', height=580)\n"
        "    return _save_html(fig, 'budget_vs_revenue_interactive', out_dir)"
    )

    heading(doc, "6.2 Top Movies by Popularity – Interactive Bar", level=2)
    body(doc,
        "Add a horizontal bar chart showing the top 10 movies by TMDB popularity "
        "score. Use px.bar() with orientation='h' and include revenue and rating "
        "in the hover tooltip:")
    code_block(doc,
        "def interactive_top_movies_bar(df, n=10, out_dir=INTERACTIVE_OUT):\n"
        "    top = df.nlargest(n, 'popularity')[\n"
        "        ['title', 'popularity', 'revenue_usd',\n"
        "         'vote_average', 'release_year', 'primary_genre']].copy()\n"
        "    top['revenue_M'] = (top['revenue_usd'] / 1e6).round(1)\n\n"
        "    fig = px.bar(\n"
        "        top.sort_values('popularity', ascending=True),\n"
        "        x='popularity', y='title', orientation='h',\n"
        "        color='primary_genre',\n"
        "        hover_name='title',\n"
        "        hover_data={'release_year': True,\n"
        "                    'vote_average': ':.2f', 'revenue_M': ':.1f'},\n"
        "        title=f'Top {n} Movies by TMDB Popularity Score',\n"
        "        template=TEMPLATE,\n"
        "    )\n"
        "    return _save_html(fig, 'top_movies_popularity_bar', out_dir)"
    )

    heading(doc, "6.3 Movies per Year – Interactive Line Chart", level=2)
    body(doc,
        "Add a line chart showing how many movies were released each year. "
        "Group the dataset by release_year and include average rating and total "
        "revenue in the hover data:")
    code_block(doc,
        "def interactive_movies_per_year(df, out_dir=INTERACTIVE_OUT):\n"
        "    yearly = (df.query('release_year >= 1980 and release_year <= 2025')\n"
        "               .groupby('release_year')\n"
        "               .agg(movie_count=('title', 'count'),\n"
        "                    avg_rating=('vote_average', 'mean'),\n"
        "                    total_revenue_B=('revenue_usd',\n"
        "                        lambda s: (s.sum() / 1e9).round(2)))\n"
        "               .reset_index())\n\n"
        "    fig = px.line(yearly, x='release_year', y='movie_count',\n"
        "                  markers=True,\n"
        "                  hover_data={'avg_rating': ':.2f',\n"
        "                              'total_revenue_B': ':.2f'},\n"
        "                  title='Movies Released per Year (1980–2025)',\n"
        "                  template=TEMPLATE)\n"
        "    fig.update_traces(line_color='#1a6faf', line_width=2.5)\n"
        "    return _save_html(fig, 'movies_per_year_line', out_dir)"
    )

    heading(doc, "6.4 Genre Ratings – Interactive Box Plot", level=2)
    body(doc,
        "Add an interactive box plot using px.box(). Pass category_orders to "
        "sort genres by median rating. Each data point shows the movie title "
        "and release year when hovered:")
    code_block(doc,
        "def interactive_genre_boxplot(df, out_dir=INTERACTIVE_OUT):\n"
        "    order = (df.groupby('primary_genre')['vote_average']\n"
        "               .median().sort_values(ascending=False).index.tolist())\n"
        "    fig = px.box(\n"
        "        df, x='primary_genre', y='vote_average',\n"
        "        category_orders={'primary_genre': order},\n"
        "        color='primary_genre',\n"
        "        hover_name='title',\n"
        "        hover_data={'release_year': True, 'vote_average': ':.2f'},\n"
        "        title='Movie Rating Distribution by Genre',\n"
        "        template=TEMPLATE,\n"
        "    )\n"
        "    fig.update_layout(showlegend=False, height=500)\n"
        "    return _save_html(fig, 'genre_rating_boxplot_interactive', out_dir)"
    )
    body(doc,
        "In the notebook, call each of these four functions and use fig.show() "
        "to display the chart inline. Hover over individual points to see "
        "the tooltip, click genre names in the legend to hide or show them, "
        "and use the box-select tool to zoom in. Open the saved HTML files "
        "in a browser to confirm they work as standalone documents.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 7 – INTERACTIVE MULTI-LAYOUT DASHBOARD
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 7 – Interactive Multi-Layout Dashboard", level=1)
    body(doc,
        "Add the final function to src/visualization/interactive_charts.py. "
        "This function builds a 2×2 interactive dashboard using "
        "plotly.graph_objects and make_subplots(). Unlike Plotly Express, "
        "Graph Objects are used here because make_subplots() requires you "
        "to add each trace manually with fig.add_trace(row=, col=):")
    code_block(doc,
        "def interactive_multi_layout(df, out_dir=INTERACTIVE_OUT):\n"
        "    fig = make_subplots(\n"
        "        rows=2, cols=2,\n"
        "        subplot_titles=('Top 10 Movies by Revenue', 'Rating Distribution',\n"
        "                        'Movies per Genre', 'Budget vs Revenue'),\n"
        "        vertical_spacing=0.14, horizontal_spacing=0.10,\n"
        "    )\n\n"
        "    # Panel 1 – top 10 revenue bar\n"
        "    top10 = (df[df['revenue_usd'] > 0].nlargest(10, 'revenue_usd')\n"
        "             .sort_values('revenue_usd'))\n"
        "    fig.add_trace(go.Bar(\n"
        "        x=top10['revenue_usd'] / 1e9, y=top10['title'],\n"
        "        orientation='h', marker_color='#1a6faf', name='Revenue',\n"
        "        hovertemplate='%{y}<br>$%{x:.2f}B<extra></extra>'),\n"
        "        row=1, col=1)\n\n"
        "    # Panel 2 – rating histogram\n"
        "    fig.add_trace(go.Histogram(\n"
        "        x=df['vote_average'].dropna(), nbinsx=20,\n"
        "        marker_color='#2ca02c', name='Ratings'),\n"
        "        row=1, col=2)\n\n"
        "    # Panel 3 – genre count bar\n"
        "    gc = df['primary_genre'].value_counts()\n"
        "    fig.add_trace(go.Bar(\n"
        "        x=gc.index, y=gc.values,\n"
        "        marker_color='#ff7f0e', name='Genre count'),\n"
        "        row=2, col=1)\n\n"
        "    # Panel 4 – budget vs revenue scatter coloured by rating\n"
        "    sd = df[(df['budget_usd'] > 0) & (df['revenue_usd'] > 0)]\n"
        "    fig.add_trace(go.Scatter(\n"
        "        x=sd['budget_usd'] / 1e6, y=sd['revenue_usd'] / 1e6,\n"
        "        mode='markers',\n"
        "        marker=dict(color=sd['vote_average'], colorscale='Viridis',\n"
        "                    showscale=True, size=8, opacity=0.7),\n"
        "        text=sd['title'],\n"
        "        hovertemplate='%{text}<br>Budget: $%{x:.0f}M<br>'\n"
        "                      'Revenue: $%{y:.0f}M<extra></extra>'),\n"
        "        row=2, col=2)\n\n"
        "    fig.update_layout(\n"
        "        title_text='Movie Industry Analytics Dashboard',\n"
        "        template=TEMPLATE, height=700, width=1100, showlegend=False)\n"
        "    return _save_html(fig, 'interactive_dashboard', out_dir)"
    )
    body(doc,
        "In the notebook, call interactive_multi_layout(df) and display the "
        "result. You should see four connected panels. Each panel is individually "
        "zoomable and hoverable. After this, interactive_charts.py is complete "
        "with 5 functions.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 8 – AUTOMATED CHART GENERATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 8 – Automated Chart Generation", level=1)
    body(doc,
        "Now create two more files that wire everything together into a "
        "single runnable command.")

    heading(doc, "8.1 Creating chart_generator.py", level=2)
    body(doc,
        "Create src/visualization/chart_generator.py. This module is the "
        "orchestrator: it loads the cleaned dataset, iterates over all static "
        "and interactive chart functions, calls each one, and prints progress. "
        "It imports from the two modules you just created:")
    code_block(doc,
        "# src/visualization/chart_generator.py\n"
        "import pandas as pd\n"
        "from pathlib import Path\n\n"
        "DATA_PATH    = Path('data/processed/cleaned/movies_clean.csv')\n"
        "STATIC_OUT   = Path('outputs/visualizations/static')\n"
        "INTERACTIVE_OUT = Path('outputs/visualizations/interactive')\n\n"
        "def load_data(path=DATA_PATH):\n"
        "    df = pd.read_csv(path, low_memory=False)\n"
        "    required = ['title', 'release_year', 'primary_genre',\n"
        "                'budget_usd', 'revenue_usd', 'vote_average',\n"
        "                'vote_count', 'popularity']\n"
        "    missing = [c for c in required if c not in df.columns]\n"
        "    if missing:\n"
        "        raise ValueError(f'Missing columns: {missing}')\n"
        "    return df\n\n"
        "def run_static_charts(df):\n"
        "    from .static_charts import (\n"
        "        plot_top_movies_by_revenue, plot_avg_rating_over_years,\n"
        "        plot_budget_vs_revenue_scatter, plot_rating_distribution,\n"
        "        plot_genre_rating_boxplot, plot_correlation_heatmap,\n"
        "        plot_genre_count_bar, plot_dashboard_subplots,\n"
        "    )\n"
        "    charts = [\n"
        "        ('top_movies_revenue',    plot_top_movies_by_revenue),\n"
        "        ('avg_rating_over_years', plot_avg_rating_over_years),\n"
        "        ('budget_vs_revenue',     plot_budget_vs_revenue_scatter),\n"
        "        ('rating_distribution',   plot_rating_distribution),\n"
        "        ('genre_rating_boxplot',  plot_genre_rating_boxplot),\n"
        "        ('correlation_heatmap',   plot_correlation_heatmap),\n"
        "        ('genre_count_bar',       plot_genre_count_bar),\n"
        "        ('dashboard_subplots',    plot_dashboard_subplots),\n"
        "    ]\n"
        "    results = {}\n"
        "    for name, fn in charts:\n"
        "        paths = fn(df, out_dir=STATIC_OUT)\n"
        "        results[name] = paths\n"
        "        print(f'  [static]  {name}')\n"
        "        print(f'            PNG -> {paths[\"png\"]}')\n"
        "    return results\n\n"
        "def generate_all(data_path=DATA_PATH):\n"
        "    df = load_data(data_path)\n"
        "    static = run_static_charts(df)\n"
        "    interactive = run_interactive_charts(df)\n"
        "    return {'static': static, 'interactive': interactive}"
    )

    heading(doc, "8.2 Creating the CLI Script", level=2)
    body(doc,
        "Create scripts/generate_visualizations.py. This is the entry point "
        "that a student runs from the project root. It uses argparse so the "
        "data path can be overridden with a --data flag:")
    code_block(doc,
        "#!/usr/bin/env python3\n"
        "# scripts/generate_visualizations.py\n"
        "import argparse, sys, os\n"
        "from pathlib import Path\n\n"
        "ROOT = Path(__file__).resolve().parents[1]\n"
        "sys.path.insert(0, str(ROOT))\n"
        "sys.path.insert(0, str(ROOT / 'src'))\n\n"
        "def main():\n"
        "    parser = argparse.ArgumentParser()\n"
        "    parser.add_argument(\n"
        "        '--data',\n"
        "        default=str(ROOT / 'data/processed/cleaned/movies_clean.csv'))\n"
        "    args = parser.parse_args()\n"
        "    os.chdir(ROOT)\n"
        "    from visualization.chart_generator import generate_all\n"
        "    generate_all(data_path=Path(args.data))\n\n"
        "if __name__ == '__main__':\n"
        "    main()"
    )

    heading(doc, "8.3 Running the Script", level=2)
    body(doc, "From the project root, run:")
    code_block(doc,
        "python scripts/generate_visualizations.py\n\n"
        "# With a custom data file:\n"
        "python scripts/generate_visualizations.py --data path/to/movies_clean.csv"
    )
    body(doc, "Expected console output:")
    code_block(doc,
        "========================================\n"
        "  Lab 12 – Data Visualization Generator\n"
        "========================================\n\n"
        "Dataset: 126 movies, 15 columns\n\n"
        "── Static charts (matplotlib / seaborn) ──\n"
        "  [static]  top_movies_revenue\n"
        "            PNG -> outputs/visualizations/static/top_movies_by_revenue.png\n"
        "            PDF -> outputs/visualizations/static/top_movies_by_revenue.pdf\n"
        "  ... (8 total)\n\n"
        "── Interactive charts (Plotly Express) ──\n"
        "  [interactive]  budget_vs_revenue\n"
        "       HTML -> outputs/visualizations/interactive/budget_vs_revenue_interactive.html\n"
        "  ... (5 total)\n\n"
        "  Done!  8 static + 5 interactive"
    )
    body(doc,
        "After running the script you should find 16 files (8 PNG + 8 PDF) "
        "in outputs/visualizations/static/ and 5 HTML files in "
        "outputs/visualizations/interactive/. Open any HTML file directly "
        "in a browser to interact with the chart.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 9 – INTEGRATE INTO RUN_PIPELINE
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 9 – Integrate into run_pipeline.py", level=1)
    body(doc,
        "Update the existing src/pipeline/run_pipeline.py to include a "
        "visualization step that runs after the analytics step from Lab 10. "
        "Add the following function to run_pipeline.py:")
    code_block(doc,
        "def run_visualizations_pipeline(cleaned_path=None):\n"
        "    \"\"\"Lab 12 – generate all static and interactive visualizations.\"\"\"\n"
        "    import sys, os\n"
        "    from pathlib import Path\n"
        "    BASE_DIR = Path(__file__).resolve().parents[2]\n"
        "    if cleaned_path is None:\n"
        "        cleaned_path = BASE_DIR / 'data/processed/cleaned/movies_clean.csv'\n"
        "    if str(BASE_DIR / 'src') not in sys.path:\n"
        "        sys.path.insert(0, str(BASE_DIR / 'src'))\n"
        "    os.chdir(BASE_DIR)\n"
        "    from visualization.chart_generator import generate_all\n"
        "    results = generate_all(data_path=Path(cleaned_path))\n"
        "    return results"
    )
    body(doc,
        "You can then call run_visualizations_pipeline() at the bottom of "
        "run_pipeline.py after run_analytics_pipeline() to make the "
        "visualization step part of the full pipeline run.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # PART 10 – DOCUMENT VISUALIZATION CHOICES
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Part 10 – Document Visualization Choices", level=1)
    body(doc,
        "Add a final section to your notebook (and include it in your "
        "submission document) that explains why each chart type was chosen "
        "for each analysis question. This is a required part of the assignment. "
        "For each chart, explain: what question it answers, why that chart type "
        "is the correct choice for that data, and what encoding decisions were "
        "made (colour, size, axis).")
    body(doc,
        "The guiding principle is to match the chart type to the statistical "
        "nature of the data: use bars for comparing discrete categories, "
        "lines for trends over time, scatter plots for relationships between "
        "two continuous variables, histograms for distributions, box plots "
        "for comparing distributions across groups, and heatmaps for "
        "pairwise correlation matrices.")

    add_table(doc,
        ["Chart", "Why this chart type"],
        [
            ["Top 10 by Revenue (horizontal bar)",
             "Ranking across named categories. Horizontal bars suit long title labels."],
            ["Rating Over Years (dual-axis line+bar)",
             "Two differently scaled time series sharing one x-axis. Bars for count, line for average."],
            ["Budget vs Revenue (scatter)",
             "Relationship between two continuous variables. Hue adds genre, size adds rating as a third dimension."],
            ["Rating Distribution (histogram+KDE)",
             "Distribution shape of a single continuous variable. KDE smooths bin artefacts."],
            ["Genre Rating Boxplot",
             "Comparing distributions across groups. Shows median, IQR, and outliers simultaneously."],
            ["Correlation Heatmap",
             "N×N pairwise statistics. Diverging colour encodes sign; annotation encodes value."],
            ["Genre Count Bar",
             "Frequency of discrete categories. Count as bar length is the most perceptually accurate channel."],
            ["2×2 Dashboard",
             "Multiple questions in one shareable figure following natural reading order."],
            ["Interactive Scatter (Plotly)",
             "Hover tooltips answer 'which specific film is this outlier?' without label clutter."],
            ["Interactive Line (Plotly)",
             "Zoom into specific decades; hover shows exact count and rating for any year."],
        ]
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # ASSIGNMENT TASK
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Assignment Task", level=1)

    heading(doc, "1. Visualization Package (25%)", level=2)
    for item in [
        "Create the src/visualization/ package with all required modules: "
        "__init__.py, static_charts.py, interactive_charts.py, chart_generator.py",
        "Implement at least 4 static charts using matplotlib and seaborn, "
        "saved as PNG (300 dpi) and PDF",
        "Implement at least 3 interactive charts using Plotly Express saved "
        "as HTML files",
        "Include a 2×2 multi-panel layout using plt.subplots(2, 2) (static) "
        "and make_subplots(rows=2, cols=2) (interactive)",
        "All chart functions must accept a DataFrame and an out_dir parameter",
    ]:
        bullet(doc, item)

    heading(doc, "2. Automated Generation (20%)", level=2)
    for item in [
        "Create src/visualization/chart_generator.py with load_data(), "
        "run_static_charts(), run_interactive_charts(), and generate_all() functions",
        "Create scripts/generate_visualizations.py as a CLI entry point with "
        "an optional --data argument",
        "Running the script from the project root must generate all charts "
        "without any manual intervention",
    ]:
        bullet(doc, item)

    heading(doc, "3. Notebook Demonstration (25%)", level=2)
    for item in [
        "Demonstrate all chart functions in notebook/lab12_visualization.ipynb",
        "Each chart must be shown inline in the notebook using fig.show() or plt.show()",
        "Include Markdown cells explaining what each chart shows and what "
        "insights you observe from your dataset",
        "The automated generator (generate_all()) must be called and its "
        "output displayed in the notebook",
    ]:
        bullet(doc, item)

    heading(doc, "4. Pipeline Integration (10%)", level=2)
    for item in [
        "Update src/pipeline/run_pipeline.py to include run_visualizations_pipeline()",
        "The function must call generate_all() and return the results dict",
    ]:
        bullet(doc, item)

    heading(doc, "5. Document Visualization Choices (20%)", level=2)
    for item in [
        "Add a 'Document Visualization Choices' section to the notebook and "
        "to your lab report",
        "For each chart, explain: what question it answers, why that chart type "
        "was chosen, and what encoding decisions were made",
        "The explanation must reference the statistical nature of the data "
        "(e.g., continuous vs. categorical, distribution vs. comparison)",
    ]:
        bullet(doc, item)

    doc.add_paragraph()

    body(doc,
        "Before running any code in this lab, make sure the following "
        "are available on your machine:")
    for req in [
        "Python 3.10 or newer  (check with: python --version)",
        "pip package manager   (comes with Python by default)",
        "The project repository cloned locally",
        "The cleaned dataset produced in Lab 9 must exist at: "
        "data/processed/cleaned/movies_clean.csv",
    ]:
        bullet(doc, req)
    doc.add_paragraph()

    # ── 2.2 Setting up a virtual environment ──
    heading(doc, "2.2 Setting Up a Virtual Environment (Recommended)", level=2)
    body(doc,
        "It is strongly recommended to work inside a virtual environment so "
        "that lab dependencies do not conflict with other Python projects "
        "on your machine.")

    body(doc, "Step 1 – Create the virtual environment (run once, from the project root):")
    code_block(doc,
        "# macOS / Linux\n"
        "python3 -m venv .venv\n\n"
        "# Windows\n"
        "python -m venv .venv"
    )

    body(doc, "Step 2 – Activate the virtual environment (run every time you open a new terminal):")
    code_block(doc,
        "# macOS / Linux\n"
        "source .venv/bin/activate\n\n"
        "# Windows (Command Prompt)\n"
        ".venv\\Scripts\\activate.bat\n\n"
        "# Windows (PowerShell)\n"
        ".venv\\Scripts\\Activate.ps1"
    )
    body(doc,
        "When activated, your terminal prompt will be prefixed with (.venv), "
        "confirming you are inside the virtual environment.")
    doc.add_paragraph()

    # ── 2.3 Installing dependencies ──
    heading(doc, "2.3 Installing All Dependencies", level=2)
    body(doc,
        "The project provides a requirements.txt file that lists every "
        "package needed across all labs. Install everything with a single command:")
    code_block(doc,
        "pip install -r requirements.txt"
    )
    body(doc,
        "If you only want to install the visualization-specific packages "
        "for this lab (faster, minimal install):")
    code_block(doc,
        "pip install matplotlib seaborn plotly kaleido pandas numpy"
    )
    body(doc,
        "To install specific minimum versions matching the lab requirements:")
    code_block(doc,
        "pip install \"matplotlib>=3.8\" \"seaborn>=0.12\" \"plotly>=6.0\" "
        "\"kaleido>=0.2\" \"pandas>=2.1\" \"numpy>=1.26\""
    )
    doc.add_paragraph()

    # ── 2.4 Verifying the installation ──
    heading(doc, "2.4 Verifying the Installation", level=2)
    body(doc,
        "After installation, confirm all packages imported correctly by "
        "running this one-liner in the terminal:")
    code_block(doc,
        "python -c \"import matplotlib, seaborn, plotly, pandas, numpy; \\\n"
        "print('matplotlib', matplotlib.__version__); \\\n"
        "print('seaborn', seaborn.__version__); \\\n"
        "print('plotly', plotly.__version__); \\\n"
        "print('pandas', pandas.__version__); \\\n"
        "print('numpy', numpy.__version__)\""
    )
    body(doc, "Expected output (versions may be higher, that is fine):")
    code_block(doc,
        "matplotlib 3.10.8\n"
        "seaborn    0.13.2\n"
        "plotly     6.7.0\n"
        "pandas     2.2.3\n"
        "numpy      1.26.4"
    )
    body(doc,
        "If you see ImportError for any package, re-run the pip install "
        "command above and ensure your virtual environment is activated.")
    doc.add_paragraph()

    # ── 2.5 Jupyter Notebook setup ──
    heading(doc, "2.5 Jupyter Notebook Setup", level=2)
    body(doc,
        "The lab includes an interactive Jupyter notebook "
        "(notebook/lab12_visualization.ipynb). To open it:")
    code_block(doc,
        "# Install Jupyter if not already available\n"
        "pip install jupyter\n\n"
        "# Launch JupyterLab from the project root\n"
        "jupyter lab\n\n"
        "# Or launch classic Notebook interface\n"
        "jupyter notebook"
    )
    body(doc,
        "In the browser that opens, navigate to the notebook/ folder "
        "and open lab12_visualization.ipynb. "
        "Run cells from top to bottom using Shift+Enter.")
    body(doc,
        "Important: the notebook uses relative paths (e.g. ../data/...), "
        "so always launch Jupyter from the project root directory, "
        "not from inside the notebook/ folder.")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 3. MATPLOTLIB – EXPLANATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "3. Matplotlib – Explanation", level=1)
    body(doc,
        "Matplotlib is Python's foundational 2-D plotting library. "
        "It follows a two-layer architecture:")

    heading(doc, "3.1 Object-Oriented API", level=2)
    body(doc,
        "Every chart begins with the explicit creation of a Figure and one "
        "or more Axes objects. This makes every visual element addressable "
        "and configurable programmatically:")
    code_block(doc,
        "fig, ax = plt.subplots(figsize=(10, 6))\n"
        "ax.barh(titles, revenues, color=colors)\n"
        "ax.set_xlabel('Revenue (USD Billion)')\n"
        "ax.set_title('Top 10 Movies by Revenue')\n"
        "fig.tight_layout()\n"
        "fig.savefig('chart.png', dpi=300, bbox_inches='tight')"
    )
    body(doc,
        "The Figure is the top-level container (the whole canvas). "
        "Each Axes is an individual plot area. Calling plt.subplots(nrows, ncols) "
        "returns a Figure and an array of Axes for multi-panel layouts.")

    heading(doc, "3.2 Saving Figures", level=2)
    body(doc,
        "fig.savefig() writes the figure to disk. Lab 12 saves every static "
        "chart in two formats:")
    for item in [
        "PNG at 300 dpi – suitable for presentations and reports.",
        "PDF (vector) – resolution-independent, ideal for print.",
    ]:
        bullet(doc, item)
    code_block(doc,
        "fig.savefig(png_path, dpi=300, bbox_inches='tight')\n"
        "fig.savefig(pdf_path, bbox_inches='tight')"
    )

    heading(doc, "3.3 Multi-Panel Layouts", level=2)
    body(doc,
        "A 2×2 dashboard is created with a single subplots() call. "
        "Each cell is addressed by its row-column index:")
    code_block(doc,
        "fig, axes = plt.subplots(2, 2, figsize=(14, 10))\n"
        "axes[0, 0].barh(...)   # top-left\n"
        "axes[0, 1].hist(...)   # top-right\n"
        "axes[1, 0].bar(...)    # bottom-left\n"
        "axes[1, 1].scatter(...)  # bottom-right\n"
        "fig.tight_layout()"
    )

    heading(doc, "3.4 Non-Interactive Backend", level=2)
    body(doc,
        "Scripts running outside a notebook must not open a display window. "
        "The Agg backend is activated at import time to enable headless rendering:")
    code_block(doc,
        "import matplotlib\n"
        "matplotlib.use('Agg')  # must be called before importing pyplot\n"
        "import matplotlib.pyplot as plt"
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 4. SEABORN – EXPLANATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "4. Seaborn – Explanation", level=1)
    body(doc,
        "Seaborn is a high-level statistical visualization library built on "
        "matplotlib. It provides publication-quality defaults, built-in "
        "color palettes, and specialized chart types for statistical data "
        "exploration with far less boilerplate.")

    heading(doc, "4.1 Theme and Context", level=2)
    code_block(doc,
        "import seaborn as sns\n\n"
        "sns.set_theme(style='whitegrid')  # axes style\n"
        "sns.set_context('notebook')       # font and element scaling\n"
        "sns.set_palette('viridis')        # default color palette"
    )
    body(doc, "Common styles: darkgrid, whitegrid, dark, white, ticks.")
    body(doc, "Common contexts: paper, notebook, talk, poster (scale up font sizes).")

    heading(doc, "4.2 Key Chart Functions", level=2)
    add_table(doc,
        ["Function", "Chart Type", "Used For"],
        [
            ["sns.histplot()", "Histogram + KDE", "Rating distribution"],
            ["sns.boxplot()", "Box-and-whisker", "Rating spread by genre"],
            ["sns.scatterplot()", "Scatter", "Budget vs Revenue with hue/size"],
            ["sns.heatmap()", "Correlation heatmap", "Numeric feature correlations"],
        ]
    )
    doc.add_paragraph()

    heading(doc, "4.3 Axes-Level vs Figure-Level", level=2)
    body(doc,
        "Seaborn provides two types of functions. Axes-level functions (e.g., "
        "histplot, boxplot) draw onto an existing Axes object passed via the "
        "ax= parameter and integrate seamlessly with matplotlib's object-oriented "
        "API. Figure-level functions (e.g., displot, catplot) create their own "
        "Figure and are less composable with multi-panel layouts.")

    heading(doc, "4.4 Passing hue= for color mapping", level=2)
    body(doc,
        "Since seaborn 0.12, passing a palette without a hue parameter is "
        "deprecated. The correct pattern for boxplots and similar charts is to "
        "use hue= with the same column as x=, and set legend=False to suppress "
        "the redundant legend:")
    code_block(doc,
        "sns.boxplot(\n"
        "    data=df, x='primary_genre', y='vote_average',\n"
        "    order=order,\n"
        "    hue='primary_genre', palette='viridis',\n"
        "    legend=False, ax=ax\n"
        ")"
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 5. PLOTLY – EXPLANATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "5. Plotly – Explanation", level=1)
    body(doc,
        "Plotly is a declarative charting library that renders figures as "
        "interactive JavaScript visualizations. It offers two APIs:")

    heading(doc, "5.1 Plotly Express", level=2)
    body(doc,
        "plotly.express (px) provides a high-level, one-function-per-chart-type "
        "interface. A complete interactive scatter plot with color, size, and "
        "hover tooltips is created in a few lines:")
    code_block(doc,
        "import plotly.express as px\n\n"
        "fig = px.scatter(\n"
        "    data,\n"
        "    x='budget_M', y='revenue_M',\n"
        "    color='primary_genre',\n"
        "    size='vote_count',\n"
        "    hover_name='title',\n"
        "    hover_data={'release_year': True, 'vote_average': ':.2f'},\n"
        "    title='Budget vs Revenue – Interactive Explorer',\n"
        "    template='plotly_white',\n"
        ")\n"
        "fig.write_html('outputs/visualizations/interactive/budget_vs_revenue.html')"
    )

    heading(doc, "5.2 Plotly Graph Objects", level=2)
    body(doc,
        "plotly.graph_objects (go) is the low-level API. It is used when "
        "finer control is needed—adding individual traces, customizing marker "
        "properties, or building multi-panel layouts with make_subplots():")
    code_block(doc,
        "import plotly.graph_objects as go\n"
        "from plotly.subplots import make_subplots\n\n"
        "fig = make_subplots(rows=2, cols=2,\n"
        "    subplot_titles=('Revenue', 'Ratings', 'Genres', 'Budget vs Revenue'))\n\n"
        "fig.add_trace(go.Bar(x=revenues, y=titles, orientation='h'), row=1, col=1)\n"
        "fig.add_trace(go.Histogram(x=ratings, nbinsx=20), row=1, col=2)\n"
        "fig.update_layout(height=700, template='plotly_white')\n"
        "fig.write_html('interactive_dashboard.html')"
    )

    heading(doc, "5.3 Saving Interactive Charts", level=2)
    body(doc,
        "Interactive charts are saved as self-contained HTML files using "
        "fig.write_html(). The file embeds all JavaScript and data, so it "
        "can be opened in any browser without an internet connection or a "
        "running Python server. For static image export, kaleido is required:")
    code_block(doc,
        "# Interactive HTML\n"
        "fig.write_html('chart.html')\n\n"
        "# Static PNG (requires kaleido)\n"
        "fig.write_image('chart.png', width=1200, height=700, scale=2)"
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 6. DATASET DESCRIPTION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "6. Dataset Description", level=1)
    body(doc,
        "The visualization module consumes the cleaned TMDB dataset produced "
        "in Lab 9. The file is located at:")
    code_block(doc, "data/processed/cleaned/movies_clean.csv")
    add_table(doc,
        ["Column", "Type", "Description"],
        [
            ["title",          "str",   "Movie title"],
            ["release_year",   "int",   "Year of theatrical release"],
            ["primary_genre",  "str",   "First genre tag (e.g., Action, Drama)"],
            ["budget_usd",     "float", "Production budget in US dollars"],
            ["revenue_usd",    "float", "Box-office revenue in US dollars"],
            ["vote_average",   "float", "TMDB audience rating 0–10"],
            ["vote_count",     "int",   "Number of votes on TMDB"],
            ["popularity",     "float", "TMDB composite popularity score"],
        ]
    )
    body(doc,
        "The dataset contains 126 movies spanning multiple decades and genres. "
        "Rows with budget_usd = 0 or revenue_usd = 0 are excluded from "
        "financial charts (they represent missing data, not zero-budget films).")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 7. STEP-BY-STEP IMPLEMENTATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "7. Step-by-Step Implementation", level=1)

    # 7.1
    heading(doc, "7.1 Project Structure", level=2)
    body(doc, "The visualization code lives in its own package under src/:")
    code_block(doc,
        "src/\n"
        "└── visualization/\n"
        "    ├── __init__.py            # package; re-exports all chart functions\n"
        "    ├── static_charts.py       # 8 matplotlib / seaborn chart functions\n"
        "    ├── interactive_charts.py  # 5 Plotly chart functions\n"
        "    └── chart_generator.py     # orchestrator: load → static → interactive\n\n"
        "scripts/\n"
        "└── generate_visualizations.py # CLI entry-point\n\n"
        "notebook/\n"
        "└── lab12_visualization.ipynb  # interactive Jupyter walkthrough\n\n"
        "outputs/visualizations/\n"
        "├── static/                    # PNG + PDF files\n"
        "└── interactive/               # HTML files"
    )

    # 7.2
    heading(doc, "7.2 Static Chart 1 – Top 10 Movies by Revenue (Bar Chart)", level=2)
    body(doc,
        "A horizontal bar chart shows the ten highest-grossing films. "
        "Revenue is displayed in USD billions with dollar-formatted tick labels "
        "and inline value annotations:")
    code_block(doc,
        "def plot_top_movies_by_revenue(df, n=10, out_dir=STATIC_OUT):\n"
        "    top = (df[df['revenue_usd'] > 0]\n"
        "           .nlargest(n, 'revenue_usd')[['title', 'revenue_usd']]\n"
        "           .sort_values('revenue_usd'))\n\n"
        "    fig, ax = plt.subplots(figsize=(10, 6))\n"
        "    colors = sns.color_palette('viridis', n)\n"
        "    bars = ax.barh(top['title'], top['revenue_usd'] / 1e9, color=colors)\n\n"
        "    for bar in bars:\n"
        "        ax.text(bar.get_width() + 0.03,\n"
        "                bar.get_y() + bar.get_height() / 2,\n"
        "                f'${bar.get_width():.2f}B', va='center', fontsize=9)\n\n"
        "    ax.set_xlabel('Box-Office Revenue (USD Billion)')\n"
        "    ax.set_title(f'Top {n} Movies by Revenue', fontweight='bold')\n"
        "    ax.xaxis.set_major_formatter(FuncFormatter(lambda x, _: f'${x:.1f}B'))\n"
        "    fig.tight_layout()\n"
        "    return _save(fig, 'top_movies_by_revenue', out_dir)"
    )

    # 7.3
    heading(doc, "7.3 Static Chart 2 – Average Rating Over Years (Dual-Axis Line)", level=2)
    body(doc,
        "A dual-axis chart overlays bar counts (number of movies per year) "
        "with a line (mean vote_average). Two separate y-axes are required "
        "because the scales differ:")
    code_block(doc,
        "fig, ax1 = plt.subplots(figsize=(12, 5))\n"
        "ax1.bar(yearly['release_year'], yearly['movie_count'],\n"
        "        color='#a8d5e2', alpha=0.5, label='Movie Count')\n"
        "ax1.set_ylabel('Number of Movies')\n\n"
        "ax2 = ax1.twinx()   # share the x-axis, independent y-axis\n"
        "ax2.plot(yearly['release_year'], yearly['avg_rating'],\n"
        "         color='#1a6faf', linewidth=2.5, marker='o')\n"
        "ax2.set_ylabel('Average Vote Rating (0–10)')\n"
        "ax2.set_ylim(0, 10)"
    )

    # 7.4
    heading(doc, "7.4 Static Chart 3 – Budget vs Revenue Scatter (seaborn)", level=2)
    body(doc,
        "A seaborn scatterplot encodes genre as hue and vote_average as "
        "marker size, adding a diagonal break-even reference line:")
    code_block(doc,
        "sns.scatterplot(\n"
        "    data=data, x='budget_M', y='revenue_M',\n"
        "    hue='genre_label', palette='viridis',\n"
        "    size='vote_average', sizes=(30, 250),\n"
        "    alpha=0.75, ax=ax\n"
        ")\n"
        "max_val = max(data['budget_M'].max(), data['revenue_M'].max()) * 1.05\n"
        "ax.plot([0, max_val], [0, max_val], '--', color='gray',\n"
        "        linewidth=1.2, label='Break-even (1:1)')"
    )

    # 7.5
    heading(doc, "7.5 Static Chart 4 – Rating Distribution (seaborn histplot + KDE)", level=2)
    code_block(doc,
        "sns.histplot(df['vote_average'].dropna(), bins=20, kde=True,\n"
        "             color='#1a6faf', edgecolor='white', ax=ax)\n"
        "ax.axvline(df['vote_average'].mean(), color='firebrick',\n"
        "           linestyle='--', label=f\"Mean = {df['vote_average'].mean():.2f}\")\n"
        "ax.axvline(df['vote_average'].median(), color='darkorange',\n"
        "           linestyle='-.', label=f\"Median = {df['vote_average'].median():.2f}\")"
    )

    # 7.6
    heading(doc, "7.6 Static Chart 5 – Genre Rating Boxplot (seaborn)", level=2)
    code_block(doc,
        "order = (df.groupby('primary_genre')['vote_average']\n"
        "           .median().sort_values(ascending=False).index.tolist())\n"
        "sns.boxplot(\n"
        "    data=df, x='primary_genre', y='vote_average',\n"
        "    order=order,\n"
        "    hue='primary_genre', palette='viridis',\n"
        "    legend=False, ax=ax\n"
        ")"
    )

    # 7.7
    heading(doc, "7.7 Static Chart 6 – Correlation Heatmap (seaborn)", level=2)
    code_block(doc,
        "numeric_cols = ['budget_usd', 'revenue_usd', 'vote_average',\n"
        "                'vote_count', 'popularity']\n"
        "corr = df[numeric_cols].corr()\n"
        "sns.heatmap(corr, annot=True, fmt='.2f', cmap='coolwarm',\n"
        "            center=0, vmin=-1, vmax=1, square=True,\n"
        "            linewidths=0.5, ax=ax)"
    )

    # 7.8
    heading(doc, "7.8 Static Chart 7 – Genre Count Bar Chart", level=2)
    code_block(doc,
        "counts = df['primary_genre'].value_counts()\n"
        "bars = ax.bar(counts.index, counts.values,\n"
        "              color=sns.color_palette('viridis', len(counts)))\n"
        "for bar in bars:\n"
        "    ax.text(bar.get_x() + bar.get_width() / 2,\n"
        "            bar.get_height() + 0.5,\n"
        "            str(int(bar.get_height())),\n"
        "            ha='center', va='bottom', fontsize=9)"
    )

    # 7.9
    heading(doc, "7.9 Static Chart 8 – 2×2 Dashboard Subplots", level=2)
    body(doc,
        "A 2×2 multi-panel figure combines the four most important views "
        "on a single canvas using matplotlib's subplot grid:")
    code_block(doc,
        "fig, axes = plt.subplots(2, 2, figsize=(14, 10))\n"
        "fig.suptitle('Movie Industry Analytics Dashboard',\n"
        "             fontsize=16, fontweight='bold')\n\n"
        "# Panel A – top 10 revenue (horizontal bar)\n"
        "axes[0, 0].barh(top['title'], top['revenue_usd'] / 1e9, ...)\n\n"
        "# Panel B – rating histogram\n"
        "sns.histplot(df['vote_average'], bins=15, kde=True, ax=axes[0, 1])\n\n"
        "# Panel C – genre counts\n"
        "axes[1, 0].bar(counts.index, counts.values, ...)\n\n"
        "# Panel D – budget vs revenue scatter\n"
        "axes[1, 1].scatter(budget, revenue, c=ratings, cmap='viridis')\n\n"
        "fig.tight_layout()"
    )

    # 7.10
    heading(doc, "7.10 Interactive Chart 1 – Budget vs Revenue (Plotly Express)", level=2)
    code_block(doc,
        "fig = px.scatter(\n"
        "    data, x='budget_M', y='revenue_M',\n"
        "    color='primary_genre',\n"
        "    size='vote_count',\n"
        "    hover_name='title',\n"
        "    hover_data={\n"
        "        'release_year': True,\n"
        "        'vote_average': ':.2f',\n"
        "        'roi': ':.1f',\n"
        "    },\n"
        "    title='Budget vs Box-Office Revenue – Interactive Explorer',\n"
        "    template='plotly_white',\n"
        ")\n"
        "# add break-even reference line\n"
        "fig.add_trace(go.Scatter(x=[0, max_v], y=[0, max_v],\n"
        "    mode='lines', line=dict(dash='dash', color='gray'),\n"
        "    name='Break-even'))\n"
        "fig.write_html('outputs/visualizations/interactive/budget_vs_revenue_interactive.html')"
    )

    # 7.11
    heading(doc, "7.11 Interactive Chart 2 – Top Movies by Popularity (px.bar)", level=2)
    code_block(doc,
        "fig = px.bar(\n"
        "    top.sort_values('popularity', ascending=True),\n"
        "    x='popularity', y='title', orientation='h',\n"
        "    color='primary_genre',\n"
        "    hover_name='title',\n"
        "    hover_data={'release_year': True, 'vote_average': ':.2f'},\n"
        "    title=f'Top {n} Movies by TMDB Popularity Score',\n"
        "    template='plotly_white',\n"
        ")"
    )

    # 7.12
    heading(doc, "7.12 Interactive Chart 3 – Movies per Year (px.line)", level=2)
    code_block(doc,
        "yearly = (df.query('release_year >= 1980 and release_year <= 2025')\n"
        "            .groupby('release_year')\n"
        "            .agg(movie_count=('title', 'count'),\n"
        "                 avg_rating=('vote_average', 'mean'))\n"
        "            .reset_index())\n\n"
        "fig = px.line(yearly, x='release_year', y='movie_count',\n"
        "              markers=True, template='plotly_white')"
    )

    # 7.13
    heading(doc, "7.13 Interactive Chart 4 – Genre Boxplot (px.box)", level=2)
    code_block(doc,
        "fig = px.box(\n"
        "    df, x='primary_genre', y='vote_average',\n"
        "    category_orders={'primary_genre': order},\n"
        "    color='primary_genre',\n"
        "    hover_name='title',\n"
        "    template='plotly_white',\n"
        ")"
    )

    # 7.14
    heading(doc, "7.14 Interactive Chart 5 – 2×2 Dashboard (Plotly Graph Objects)", level=2)
    code_block(doc,
        "fig = make_subplots(rows=2, cols=2,\n"
        "    subplot_titles=('Top 10 Revenue', 'Rating Distribution',\n"
        "                    'Movies per Genre', 'Budget vs Revenue'))\n\n"
        "fig.add_trace(go.Bar(x=revenues, y=titles, orientation='h'),  row=1, col=1)\n"
        "fig.add_trace(go.Histogram(x=ratings, nbinsx=20),              row=1, col=2)\n"
        "fig.add_trace(go.Bar(x=genre_names, y=genre_counts),           row=2, col=1)\n"
        "fig.add_trace(go.Scatter(x=budgets, y=revenues, mode='markers'), row=2, col=2)\n\n"
        "fig.update_layout(height=700, template='plotly_white',\n"
        "                  title_text='Movie Industry Analytics Dashboard')\n"
        "fig.write_html('outputs/visualizations/interactive/interactive_dashboard.html')"
    )

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 8. AUTOMATED CHART GENERATION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "8. Automated Chart Generation", level=1)
    body(doc,
        "All chart functions are invoked by the orchestrator module "
        "src/visualization/chart_generator.py, which iterates over a list "
        "of (name, function) pairs, calls each with the DataFrame and "
        "output directory, and prints progress:")
    code_block(doc,
        "charts = [\n"
        "    ('top_movies_revenue',    plot_top_movies_by_revenue),\n"
        "    ('avg_rating_over_years', plot_avg_rating_over_years),\n"
        "    ('budget_vs_revenue',     plot_budget_vs_revenue_scatter),\n"
        "    # ... etc.\n"
        "]\n"
        "for name, fn in charts:\n"
        "    paths = fn(df, out_dir=STATIC_OUT)\n"
        "    print(f'  [static]  {name}')\n"
        "    print(f'            PNG → {paths[\"png\"]}')"
    )
    body(doc, "The CLI script exposes a single entry point:")
    code_block(doc,
        "# From the project root:\n"
        "python scripts/generate_visualizations.py\n\n"
        "# Custom data path:\n"
        "python scripts/generate_visualizations.py --data path/to/movies_clean.csv"
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 9. EXPORTED FILES
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "9. Exported File Locations", level=1)
    body(doc,
        "Running the script produces 13 output files across two directories:")

    heading(doc, "9.1 Static Charts (PNG + PDF)", level=2)
    add_table(doc,
        ["File (stem)", "Chart Type", "Key Insight"],
        [
            ["top_movies_by_revenue",    "Horizontal bar",    "Revenue leaders by film"],
            ["avg_rating_over_years",    "Dual-axis line+bar","Rating trends 1980–2025"],
            ["budget_vs_revenue_scatter","Scatter",           "ROI landscape by genre"],
            ["rating_distribution",      "Histogram + KDE",   "Audience rating bell curve"],
            ["genre_rating_boxplot",     "Box-whisker",       "Rating spread per genre"],
            ["correlation_heatmap",      "Heatmap",           "Feature correlation matrix"],
            ["genre_count_bar",          "Vertical bar",      "Genre frequency"],
            ["dashboard_subplots",       "2×2 multi-panel",   "All-in-one overview"],
        ]
    )
    body(doc, "Output directory: outputs/visualizations/static/")

    heading(doc, "9.2 Interactive Charts (HTML)", level=2)
    add_table(doc,
        ["HTML File", "Chart Type"],
        [
            ["budget_vs_revenue_interactive.html", "Plotly Express scatter with hover"],
            ["top_movies_popularity_bar.html",     "Plotly Express horizontal bar"],
            ["movies_per_year_line.html",           "Plotly Express line chart"],
            ["genre_rating_boxplot_interactive.html","Plotly Express boxplot"],
            ["interactive_dashboard.html",          "Plotly make_subplots 2×2 dashboard"],
        ]
    )
    body(doc, "Output directory: outputs/visualizations/interactive/")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 10. HOW TO RUN THE SCRIPT
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "10. How to Run the Script", level=1)
    body(doc, "Prerequisites – activate the project virtual environment:")
    code_block(doc,
        "# macOS / Linux\n"
        "source .venv/bin/activate\n\n"
        "# Windows\n"
        ".venv\\Scripts\\activate"
    )
    body(doc, "Install (or verify) required packages:")
    code_block(doc, "pip install plotly kaleido seaborn matplotlib")
    body(doc, "Run the visualization generator from the project root:")
    code_block(doc, "python scripts/generate_visualizations.py")
    body(doc, "Expected console output:")
    code_block(doc,
        "========================================\n"
        "  Lab 12 – Data Visualization Generator\n"
        "========================================\n\n"
        "Dataset: 126 movies, 15 columns\n\n"
        "── Static charts (matplotlib / seaborn) ──\n"
        "  [static]  top_movies_revenue\n"
        "            PNG → outputs/visualizations/static/top_movies_by_revenue.png\n"
        "            PDF → outputs/visualizations/static/top_movies_by_revenue.pdf\n"
        "  ... (8 total)\n\n"
        "── Interactive charts (Plotly Express) ──\n"
        "  [interactive]  budget_vs_revenue\n"
        "                 HTML → outputs/visualizations/interactive/budget_vs_revenue_interactive.html\n"
        "  ... (5 total)\n\n"
        "========================================\n"
        "  Done!  8 static + 5 interactive\n"
        "========================================"
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 11. ASSIGNMENT 12 SOLUTION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "11. Assignment 12 Solution", level=1)
    body(doc,
        "The assignment required implementing a data visualization pipeline "
        "for the TMDB movie dataset with a minimum of two static and two "
        "interactive charts, plus a section documenting visualization choices.")

    heading(doc, "Deliverables Completed", level=2)
    deliverables = [
        ("src/visualization/__init__.py",         "Package initialization; re-exports all chart functions"),
        ("src/visualization/static_charts.py",    "8 static matplotlib/seaborn charts"),
        ("src/visualization/interactive_charts.py","5 interactive Plotly charts"),
        ("src/visualization/chart_generator.py",  "Orchestrator; load → static → interactive pipeline"),
        ("scripts/generate_visualizations.py",    "CLI entry-point with argparse --data flag"),
        ("notebook/lab12_visualization.ipynb",    "Jupyter notebook with step-by-step walkthrough"),
        ("outputs/visualizations/static/",        "8 charts × 2 formats = 16 files (PNG + PDF)"),
        ("outputs/visualizations/interactive/",   "5 HTML files"),
        ("requirements.txt",                      "Added plotly>=6.0.0 and kaleido>=0.2.1"),
        ("src/pipeline/run_pipeline.py",           "Added run_visualizations_pipeline() function"),
    ]
    add_table(doc, ["File / Directory", "Description"], deliverables)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 12. DOCUMENT VISUALIZATION CHOICES
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "12. Document Visualization Choices", level=1)
    body(doc,
        "Each chart type was selected to match the statistical nature of the "
        "data it represents. The guiding principles are: "
        "(1) choose the chart type that answers the specific question; "
        "(2) encode the most important variable on the most perceptual channel "
        "(position > length > color > size); "
        "(3) minimize non-data ink.")

    choices = [
        (
            "Bar chart (horizontal) – Top Movies by Revenue",
            "The question is a ranking comparison across named categories (film titles). "
            "Horizontal bars are preferred over vertical when category labels are long "
            "strings, because they read left-to-right naturally. The viridis palette "
            "encodes rank via luminance. Inline value annotations eliminate the need "
            "for the reader to scan the x-axis.",
        ),
        (
            "Dual-axis Line + Bar – Rating Over Years",
            "Two related but differently scaled time series are displayed together: "
            "movie count (count variable, bar) and average rating (continuous 0–10, "
            "line). A dual-axis chart is justified here because the two series share "
            "a temporal x-axis and a reader naturally wants to correlate them. "
            "Bars encode count (additive quantity, length maps well) while the line "
            "connects continuous yearly averages.",
        ),
        (
            "Scatter plot – Budget vs Revenue",
            "The central financial question is whether higher budgets yield higher "
            "revenue. A scatter plot is the canonical chart for two continuous "
            "variables where we want to see correlation, outliers, and clusters. "
            "Genre is encoded as hue (categorical color) to reveal whether genre "
            "moderates the budget-revenue relationship. Vote average is encoded "
            "as marker size because it is a secondary dimension. The break-even "
            "diagonal makes profitability immediately visible.",
        ),
        (
            "Histogram + KDE – Rating Distribution",
            "The question is about the shape of a single continuous variable's "
            "distribution—whether it is normal, skewed, bi-modal, etc. "
            "A histogram reveals frequency; the KDE overlay smooths bin-boundary "
            "artefacts and shows the underlying density. Mean and median reference "
            "lines highlight the central tendency and any skew.",
        ),
        (
            "Box-and-whisker – Rating by Genre",
            "When comparing distributions of a continuous variable (vote_average) "
            "across many groups (genres), box plots are more space-efficient than "
            "overlaid histograms. They show the median, interquartile range, and "
            "outliers simultaneously. Genres are sorted by median rating to make "
            "the ranking immediately readable.",
        ),
        (
            "Heatmap – Correlation Matrix",
            "A correlation matrix is an N×N grid of pairwise statistics. "
            "A heatmap is the standard encoding: the diverging coolwarm palette "
            "maps negative correlations to cool blue and positive to warm red, "
            "with white at zero. Annotating each cell with the numeric value "
            "prevents misreading from the color scale. The square aspect ratio "
            "ensures equal visual weight for all cell pairs.",
        ),
        (
            "Vertical Bar – Genre Count",
            "The question is a frequency distribution across a discrete categorical "
            "variable (genre). Bars encode count as length—the most perceptually "
            "accurate channel for magnitude. Vertical orientation is used here "
            "because genre names are short and the chart is wider than it is tall. "
            "Inline count labels above each bar eliminate axis reading.",
        ),
        (
            "2×2 Multi-panel Dashboard",
            "An executive dashboard must convey several different questions at "
            "a glance. A 2×2 grid combines the four most informative individual "
            "views into a single shareable figure. The layout follows a natural "
            "reading order: top-left (top revenue), top-right (distribution), "
            "bottom-left (genre composition), bottom-right (budget/revenue relationship).",
        ),
        (
            "Interactive Scatter (Plotly Express) – Budget vs Revenue",
            "The static scatter plot raises a follow-up question: 'Which specific "
            "film is this outlier?' Interactive hover tooltips answer it without "
            "overcrowding the chart with text labels. The Plotly version adds ROI "
            "to the tooltip so analysts can assess profitability for individual films.",
        ),
        (
            "Interactive Line – Movies per Year",
            "A time-series line chart is the correct choice when the reader wants "
            "to track a single metric's evolution. Interactivity lets users zoom "
            "into specific decades and hover to see exact counts and average ratings "
            "for any year.",
        ),
    ]

    for chart_name, rationale in choices:
        p = doc.add_paragraph()
        run_bold = p.add_run(chart_name + ": ")
        run_bold.bold = True
        run_bold.font.color.rgb = RGBColor(0x1A, 0x6F, 0xAF)
        p.add_run(rationale)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # 13. CONCLUSION
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "13. Conclusion", level=1)
    body(doc,
        "Lab 12 extended the Movie Industry Analytics Pipeline with a "
        "complete, production-ready visualization layer. The implementation "
        "demonstrates how matplotlib's object-oriented API, seaborn's "
        "statistical chart templates, and Plotly's interactive rendering "
        "complement each other: matplotlib and seaborn produce crisp, "
        "publication-quality static figures (PNG/PDF) while Plotly generates "
        "interactive HTML reports that allow stakeholders to explore the data "
        "without writing code.")
    body(doc,
        "Eight static and five interactive charts were generated automatically "
        "by a single CLI command. The charts cover the full range of common "
        "analysis patterns: ranking (bar), temporal trends (line), distributions "
        "(histogram, box), correlation (scatter, heatmap), and composition (grouped bar). "
        "Visualization choices were guided by the Tufte principle of maximizing "
        "data-ink ratio—every pixel encodes information, and chart types were "
        "matched to the statistical nature of the question being answered.")
    body(doc,
        "The modular architecture—with each chart function accepting a DataFrame "
        "and an output directory—makes it trivial to regenerate all charts after "
        "data updates, swap the dataset for a different source, or add new chart "
        "functions without touching the orchestrator. This design aligns with the "
        "broader pipeline philosophy established in earlier labs: load, transform, "
        "analyze, and now visualize—each step encapsulated, composable, and "
        "reproducible.")

    # ── save ──────────────────────────────────────────────────────────────────
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    ROOT = Path(__file__).resolve().parents[1]
    out = ROOT / "docs" / "labs" / "lab12_data_visualization.docx"
    build_document(out)
